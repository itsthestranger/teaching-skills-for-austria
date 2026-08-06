#!/usr/bin/env python3
"""``docs/`` teacher-material ingestion (plan §6.6, backlog E6-05).

This module owns the *full* ingestion contract for the teacher's personal
``docs/`` folder: the folder convention, the subject alias table, format
conversion (``.md``/``.txt`` native, ``.pdf``/``.docx`` converted to
Markdown-ish text and cached under ``docs/.cache/``), and the request-scoped
limits (max file size, max file count, an approximate token budget).

``plugin/scripts/kompetenz.py``'s ``finde_lernaufgaben`` is the only public
entry point a skill calls (plan §5); it delegates to :func:`sammle` here and
returns just the accepted entries as plain dicts, keeping its own contract
(``list[dict]``, never raising for "nothing found") unchanged. This module
additionally exposes :func:`sammle` itself for callers that want the full,
observable report -- including what was skipped or dropped and why -- the
same way ``kompetenz.stichwort_abdeckung`` exposes routing detail beyond the
nine contracted ``finde_*`` functions.

**Never fabricate a competence attribution.** A file's ``kompetenz_id`` is
only ever set from an explicit binding the file itself carries (a filename
suffix or YAML frontmatter key) -- never guessed from keywords or folder
position. Folder-derived ``fach``/``stufe`` are informational metadata, not
a competence link.

**Everything here is teacher-supplied, never official.** Every returned
entry carries ``herkunft: "docs"`` and ``amtlich: False`` -- reusing the
vocabulary ``plugin/skills/*/scripts/lesson_common.py``'s ``resolve_herkunft``
already treats as authoritative (only ``amtlich is True`` ever renders as
official RIS content). This module does not import the renderer and does not
invent a second origin mechanism.

**Nothing is silently dropped.** Oversized files, unconvertible PDFs/DOCX,
and anything trimmed by the file-count or token-budget cap are logged
(``LOGGER.warning``) *and* recorded in the returned :class:`Ingestionsbericht`
so a caller can surface "N Dateien nicht beruecksichtigt" rather than a
result that silently looks complete.

Pure stdlib plus the already-accepted runtime dependency ``python-docx``
(decision E5-11) for ``.docx`` reading. **No PDF library is a project
dependency.** PDF extraction is a small, explicit seam
(:data:`_PDF_EXTRAKTOREN`): each entry is tried in order and must already be
importable, so with nothing installed (the shipped state) every PDF is
logged as unusable and skipped -- never fatal, exactly the path §6.6
prescribes for scanned PDFs. Dropping in a future extractor means appending
one function to that tuple, nothing else.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

LOGGER = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Folder convention (plan §6.6)
# ---------------------------------------------------------------------------

#: Plain-language folder name (casefolded) -> subject code. Deliberately a
#: small, explicit table (plan §6.6), not derived from the shipped shard
#: registry: docs/ folders name a *subject* only, never a band -- the same
#: "mathematik" folder can hold material for PRIM.M and SEK1.M alike, and
#: the stufe token (SCH.. vs K..), when present, is what disambiguates.
FACH_ALIAS: dict[str, str] = {
    "mathematik": "M",
    "deutsch": "D",
    "englisch": "E",
    "sachunterricht": "SU",
    "m": "M",
    "d": "D",
    "e": "E",
    "su": "SU",
}

_STUFE_TOKENS = {f"K{n}" for n in range(1, 5)} | {f"SCH{n}" for n in range(1, 5)}

#: Filename-suffix competence binding: ``…__AT.LP23.SEK1.M.ZAHLEN.K2.03.md``.
#: Deliberately anchored on the literal ``AT.LP23.`` prefix so an incidental
#: double underscore elsewhere in a teacher's filename can never be mistaken
#: for a binding.
_ID_SUFFIX_RE = re.compile(r"^AT\.LP23\.[A-Z0-9]+(?:\.[A-Z0-9]+)+$")

_FRONTMATTER_RE = re.compile(r"^---\r?\n(.*?)\r?\n---\r?\n?", re.DOTALL)

# ---------------------------------------------------------------------------
# Limits (plan §6.6 -- "defaults, overridable")
# ---------------------------------------------------------------------------

#: Max size of a single source file that is read at all.
MAX_DATEIGROESSE_BYTES = 2 * 1024 * 1024
#: Max number of files returned for one request (post-filter, post-ranking).
MAX_DATEIEN = 20
#: Approximate cumulative token budget for the returned material.
TOKEN_BUDGET = 4000

_NATIVE_SUFFIXE = {".md", ".txt"}
_KONVERTIERBARE_SUFFIXE = {".pdf", ".docx"}

#: Default ingestion root when a caller does not name one explicitly --
#: ``docs/`` relative to the process's working directory (matches the
#: pre-E6-05 interim reader's behaviour). Module-level and therefore
#: injectable: tests point this at a committed fixture tree with
#: ``monkeypatch.setattr(docs_ingest, "STANDARD_DOCS_ROOT", ...)``, the same
#: seam ``data-pipeline/tests/test_kompetenz_contract.py`` uses for
#: ``kompetenz.KOMPETENZEN_ROOT`` -- ``docs/`` is gitignored, so a test must
#: never depend on its real contents.
STANDARD_DOCS_ROOT = Path("docs")


def approx_tokens(text: str) -> int:
    """``len(utf-8 bytes) / 4`` -- the same order-of-magnitude heuristic
    ``data-pipeline/build_dataset.py``'s ``approx_tokens`` uses for the
    §6.7 shard budget, reused here for the §6.6 per-turn ``docs/`` budget so
    the project states one token-approximation method, not two."""
    return max(1, len(text.encode("utf-8")) // 4)


# ---------------------------------------------------------------------------
# Result / report shapes
# ---------------------------------------------------------------------------


@dataclass
class Lernaufgabe:
    """One accepted teacher-supplied document."""

    titel: str
    pfad: str
    fach: str | None
    stufe: str | None
    kompetenz_id: str | None
    format: str  # "md" | "txt" | "pdf" | "docx"
    konvertiert: bool
    text: str
    bytes: int
    tokens_approx: int
    herkunft: str = "docs"
    amtlich: bool = False

    def to_dict(self) -> dict[str, Any]:
        return {
            "titel": self.titel,
            "pfad": self.pfad,
            "fach": self.fach,
            "stufe": self.stufe,
            "kompetenz_id": self.kompetenz_id,
            "format": self.format,
            "konvertiert": self.konvertiert,
            "text": self.text,
            "bytes": self.bytes,
            "tokens_approx": self.tokens_approx,
            "herkunft": self.herkunft,
            "amtlich": self.amtlich,
        }


@dataclass
class Ingestionsbericht:
    """Full, observable outcome of one :func:`sammle` call."""

    treffer: list[Lernaufgabe] = field(default_factory=list)
    #: Files that could not be read/converted at all: {"pfad", "grund"}.
    uebersprungen: list[dict[str, str]] = field(default_factory=list)
    #: Matching files trimmed by the file-count or token-budget cap:
    #: {"pfad", "grund"}. Ranked-but-not-included, never silently absent.
    verworfen: list[dict[str, str]] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Folder-convention parsing
# ---------------------------------------------------------------------------


def fach_aus_ordner(ordnername: str) -> str | None:
    """Map one top-level ``docs/<ordner>`` name to a subject code, or
    ``None`` if the folder name is not in :data:`FACH_ALIAS` -- an unknown
    subject folder, never a raised error (plan §6.6: unassigned, not
    discarded)."""
    return FACH_ALIAS.get(ordnername.strip().casefold())


def stufe_normalisieren(roh: str) -> str | None:
    kandidat = roh.strip().upper()
    return kandidat if kandidat in _STUFE_TOKENS else None


# ---------------------------------------------------------------------------
# Competence binding: filename suffix (precedence) then YAML frontmatter
# ---------------------------------------------------------------------------


def _kompetenz_id_aus_dateiname(stem: str) -> tuple[str | None, str]:
    """``("…__ID", basisname)`` -> ``(ID, basisname)`` if the suffix matches
    the binding grammar, else ``(None, stem)`` unchanged."""
    basis, sep, kandidat = stem.rpartition("__")
    if sep and _ID_SUFFIX_RE.match(kandidat):
        return kandidat, basis
    return None, stem


def _frontmatter_parsen(text: str) -> tuple[dict[str, str], str]:
    """Minimal ``key: value`` YAML-frontmatter reader (no external YAML
    dependency -- deliberately not a general parser, only flat scalar keys,
    which is all the plan's optional frontmatter binding needs). Returns
    ``(felder, restlicher_text)``; ``felder`` is ``{}`` when there is no
    frontmatter block or a line does not parse as ``key: value``."""
    m = _FRONTMATTER_RE.match(text)
    if not m:
        return {}, text
    felder: dict[str, str] = {}
    for zeile in m.group(1).splitlines():
        if ":" not in zeile:
            continue
        schluessel, _, wert = zeile.partition(":")
        schluessel = schluessel.strip().strip("\"'")
        wert = wert.strip().strip("\"'")
        if schluessel:
            felder[schluessel] = wert
    return felder, text[m.end():]


def _titel_aus_text(text: str, fallback: str) -> str:
    for zeile in text.splitlines():
        zeile = zeile.strip()
        if zeile.startswith("#"):
            return zeile.lstrip("#").strip() or fallback
        if zeile:
            # First non-empty, non-heading line: keep scanning up to the
            # first real heading, but do not treat body prose as a title.
            continue
    return fallback


# ---------------------------------------------------------------------------
# PDF extraction seam -- no library shipped or required; a future extractor
# is a new entry in this tuple, nothing else changes.
# ---------------------------------------------------------------------------


def _pdf_text_via_pypdf(pfad: Path) -> str | None:
    try:
        import pypdf  # type: ignore[import-not-found]
    except ImportError:
        return None
    reader = pypdf.PdfReader(str(pfad))
    text = "\n".join((seite.extract_text() or "") for seite in reader.pages)
    return text.strip() or None


def _pdf_text_via_pdfminer(pfad: Path) -> str | None:
    try:
        from pdfminer.high_level import extract_text  # type: ignore[import-not-found]
    except ImportError:
        return None
    text = extract_text(str(pfad)) or ""
    return text.strip() or None


#: Ordered seam of optional PDF-text extractors. Each callable either returns
#: importable-and-successful text, ``None`` (not installed / no usable
#: text, e.g. a scanned PDF), or raises -- a raise is caught by
#: :func:`extrahiere_pdf_text` and treated the same as ``None``. None of
#: these libraries is a project dependency (decision, E6-05): with nothing
#: installed, every entry returns ``None`` and PDFs are logged unusable.
_PDF_EXTRAKTOREN: tuple[Callable[[Path], str | None], ...] = (
    _pdf_text_via_pypdf,
    _pdf_text_via_pdfminer,
)


def extrahiere_pdf_text(pfad: Path) -> str | None:
    """Try every registered extractor in turn. Returns ``None`` (never
    raises) when no extractor is installed or none produced usable text --
    the caller logs that as "unusable, skipped", not fatal."""
    for extraktor in _PDF_EXTRAKTOREN:
        try:
            text = extraktor(pfad)
        except Exception as exc:  # noqa: BLE001 -- a hostile/corrupt PDF must never be fatal
            LOGGER.warning("PDF-Extraktor %s an %s gescheitert: %s", extraktor.__name__, pfad, exc)
            continue
        if text:
            return text
    return None


# ---------------------------------------------------------------------------
# DOCX conversion -- python-docx is an already-accepted runtime dependency
# (E5-11); reading is the same library the renderer already ships with.
# ---------------------------------------------------------------------------


def _konvertiere_docx(pfad: Path) -> str | None:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        LOGGER.warning("python-docx nicht installiert -- %s wird uebersprungen", pfad)
        return None
    try:
        dokument = docx.Document(str(pfad))
    except Exception as exc:  # noqa: BLE001 -- a corrupt/foreign .docx must never be fatal
        LOGGER.warning("DOCX konnte nicht gelesen werden (%s): %s", pfad, exc)
        return None

    zeilen: list[str] = []
    ueberschrift_re = re.compile(r"^Heading\s*(\d+)", re.IGNORECASE)
    for absatz in dokument.paragraphs:
        text = absatz.text.strip()
        if not text:
            continue
        stilname = (absatz.style.name if absatz.style is not None else "") or ""
        m = ueberschrift_re.match(stilname)
        if m:
            ebene = min(max(int(m.group(1)), 1), 6)
            zeilen.append(f"{'#' * ebene} {text}")
        else:
            zeilen.append(text)
    for tabelle in dokument.tables:
        for zeile in tabelle.rows:
            zellen = [zelle.text.strip() for zelle in zeile.cells]
            if any(zellen):
                zeilen.append(" | ".join(zellen))

    text = "\n\n".join(zeilen).strip()
    return text or None


# ---------------------------------------------------------------------------
# Conversion cache (docs/.cache/) -- source files are never written to.
# ---------------------------------------------------------------------------


def _cache_pfad(rel: Path, docs_root: Path) -> Path:
    return docs_root / ".cache" / rel.parent / (rel.name + ".md")


def _konvertiert_lesen_oder_erzeugen(
    quelle: Path, rel: Path, docs_root: Path, konverter: Callable[[Path], str | None]
) -> str | None:
    """Read a cached conversion if it is at least as new as the source,
    else convert and cache. Returns ``None`` (without writing a cache
    entry) if conversion failed -- retried on the next call, which is cheap
    and keeps a transient failure from becoming permanent."""
    cache = _cache_pfad(rel, docs_root)
    try:
        if cache.is_file() and cache.stat().st_mtime >= quelle.stat().st_mtime:
            return cache.read_text(encoding="utf-8")
    except OSError:
        pass  # fall through to reconversion

    text = konverter(quelle)
    if text is None:
        return None
    try:
        cache.parent.mkdir(parents=True, exist_ok=True)
        cache.write_text(text, encoding="utf-8")
    except OSError as exc:
        LOGGER.warning("Konvertierungs-Cache konnte nicht geschrieben werden (%s): %s", cache, exc)
    return text


# ---------------------------------------------------------------------------
# Relevance ranking (beyond the caps: rank, then log what is dropped)
# ---------------------------------------------------------------------------


def _relevanz_schluessel(
    eintrag: Lernaufgabe, ziel_fach: str | None, ziel_stufe: str | None, ziel_kompetenz_id: str | None
) -> tuple[int, int, int, str]:
    id_treffer = 0 if (ziel_kompetenz_id and eintrag.kompetenz_id == ziel_kompetenz_id) else 1
    fach_treffer = 0 if (ziel_fach and eintrag.fach == ziel_fach) else 1
    if ziel_stufe and eintrag.stufe == ziel_stufe:
        stufe_treffer = 0
    elif eintrag.stufe is None:
        stufe_treffer = 1  # subject-wide: relevant but less specific
    else:
        stufe_treffer = 2
    return (id_treffer, fach_treffer, stufe_treffer, eintrag.pfad)


# ---------------------------------------------------------------------------
# The walker
# ---------------------------------------------------------------------------


def sammle(
    docs_root: str | Path,
    *,
    fach: str | None = None,
    stufe: str | None = None,
    kompetenz_id: str | None = None,
    max_dateien: int = MAX_DATEIEN,
    max_dateigroesse_bytes: int = MAX_DATEIGROESSE_BYTES,
    token_budget: int = TOKEN_BUDGET,
) -> Ingestionsbericht:
    """Walk ``docs_root`` per the plan §6.6 folder convention and return a
    full :class:`Ingestionsbericht`.

    ``fach`` is a bare subject code (``"M"``/``"D"``/``"E"``/``"SU"``) or a
    ``"<BAND>.<FACH>"`` shard key (only the ``.FACH`` half is used -- a
    ``docs/`` folder never encodes a band). ``stufe`` is ``SCH1``..``SCH4``
    or ``K1``..``K4``. Missing/empty ``docs_root`` returns an empty,
    non-error report (plan §6.6).

    A file whose folder is not in :data:`FACH_ALIAS` is *unassigned*
    (``fach=None``, ``stufe=None``) rather than discarded when browsing
    without a ``fach`` filter; a caller that names a specific ``fach``
    naturally does not get unrelated unassigned material back, the same way
    it would not get a different subject's material back.
    """
    wurzel = Path(docs_root)
    bericht = Ingestionsbericht()
    if not wurzel.is_dir():
        return bericht

    ziel_fach = fach.strip().upper().rsplit(".", 1)[-1] if fach else None
    ziel_stufe = stufe.strip().upper() if stufe else None
    ziel_kompetenz_id = kompetenz_id.strip() if kompetenz_id else None

    kandidaten: list[Lernaufgabe] = []
    for pfad in sorted(wurzel.rglob("*")):
        if not pfad.is_file():
            continue
        rel = pfad.relative_to(wurzel)
        teile = rel.parts
        if any(teil.startswith(".") for teil in teile):
            # Hidden files/dirs, notably docs/.cache/ (the conversion cache
            # itself, never source content) and dotfiles in general.
            continue
        suffix = pfad.suffix.lower()
        if len(teile) == 1 and pfad.name.casefold() == "readme.md":
            # docs/README.md is this folder's own scaffolding (the one
            # non-.gitkeep file the repo commits at docs/ root), never a
            # teacher-authored Lernaufgabe.
            continue
        if suffix not in _NATIVE_SUFFIXE and suffix not in _KONVERTIERBARE_SUFFIXE:
            continue

        try:
            groesse = pfad.stat().st_size
        except OSError as exc:
            bericht.uebersprungen.append({"pfad": str(rel), "grund": f"nicht lesbar: {exc}"})
            LOGGER.warning("docs_ingest: %s nicht lesbar: %s", rel, exc)
            continue
        if groesse > max_dateigroesse_bytes:
            grund = f"Datei ueberschreitet das {max_dateigroesse_bytes}-Byte-Limit ({groesse} Byte)"
            bericht.uebersprungen.append({"pfad": str(rel), "grund": grund})
            LOGGER.warning("docs_ingest: %s uebersprungen -- %s", rel, grund)
            continue

        konvertiert = suffix in _KONVERTIERBARE_SUFFIXE
        if suffix == ".pdf":
            rohtext = _konvertiert_lesen_oder_erzeugen(pfad, rel, wurzel, extrahiere_pdf_text)
            if rohtext is None:
                grund = (
                    "PDF nicht extrahierbar (kein Extraktor installiert oder gescannt/"
                    "bildbasiert) -- uebersprungen, nicht fatal"
                )
                bericht.uebersprungen.append({"pfad": str(rel), "grund": grund})
                LOGGER.warning("docs_ingest: %s -- %s", rel, grund)
                continue
        elif suffix == ".docx":
            rohtext = _konvertiert_lesen_oder_erzeugen(pfad, rel, wurzel, _konvertiere_docx)
            if rohtext is None:
                grund = "DOCX nicht konvertierbar -- uebersprungen, nicht fatal"
                bericht.uebersprungen.append({"pfad": str(rel), "grund": grund})
                LOGGER.warning("docs_ingest: %s -- %s", rel, grund)
                continue
        else:
            try:
                rohtext = pfad.read_text(encoding="utf-8", errors="replace")
            except OSError as exc:
                bericht.uebersprungen.append({"pfad": str(rel), "grund": f"nicht lesbar: {exc}"})
                LOGGER.warning("docs_ingest: %s nicht lesbar: %s", rel, exc)
                continue

        frontmatter, koerper = _frontmatter_parsen(rohtext) if suffix in _NATIVE_SUFFIXE else ({}, rohtext)
        kompetenz_id_datei, basisstem = _kompetenz_id_aus_dateiname(pfad.stem)
        if kompetenz_id_datei is None and frontmatter.get("kompetenz_id"):
            kompetenz_id_datei = frontmatter["kompetenz_id"]

        fallback_titel = frontmatter.get("titel") or frontmatter.get("title") or basisstem
        titel = _titel_aus_text(koerper, fallback_titel)

        erkannter_fach = fach_aus_ordner(teile[0]) if teile else None
        erkannte_stufe = None
        if erkannter_fach and len(teile) >= 3:
            erkannte_stufe = stufe_normalisieren(teile[1])

        # An explicit kompetenz_id binding (filename suffix / frontmatter)
        # takes precedence over the folder location (plan §6.6) -- a file
        # correctly bound to the requested competence is never excluded
        # merely because a teacher filed it under a different subject/level
        # folder. Absent such a binding match, fall back to the folder
        # convention's subject+level filter as before.
        kompetenz_id_treffer = ziel_kompetenz_id is not None and kompetenz_id_datei == ziel_kompetenz_id
        if not kompetenz_id_treffer:
            if ziel_fach is not None and erkannter_fach != ziel_fach:
                continue
            if ziel_stufe is not None and erkannte_stufe is not None and erkannte_stufe != ziel_stufe:
                continue

        kandidaten.append(
            Lernaufgabe(
                titel=titel,
                pfad=str(rel),
                fach=erkannter_fach,
                stufe=erkannte_stufe,
                kompetenz_id=kompetenz_id_datei,
                format=suffix.lstrip("."),
                konvertiert=konvertiert,
                text=koerper.strip(),
                bytes=groesse,
                tokens_approx=approx_tokens(koerper),
            )
        )

    kandidaten.sort(key=lambda e: _relevanz_schluessel(e, ziel_fach, ziel_stufe, ziel_kompetenz_id))

    ueber_limit = kandidaten[max_dateien:]
    kandidaten = kandidaten[:max_dateien]
    for e in ueber_limit:
        grund = f"ueber dem Dateilimit ({max_dateien}) -- nach Relevanz verworfen"
        bericht.verworfen.append({"pfad": e.pfad, "grund": grund})
        LOGGER.warning("docs_ingest: %s -- %s", e.pfad, grund)

    laufendes_budget = 0
    for e in kandidaten:
        if bericht.treffer and laufendes_budget + e.tokens_approx > token_budget:
            grund = f"ueber dem Token-Budget (~{token_budget}) -- nach Relevanz verworfen"
            bericht.verworfen.append({"pfad": e.pfad, "grund": grund})
            LOGGER.warning("docs_ingest: %s -- %s", e.pfad, grund)
            continue
        laufendes_budget += e.tokens_approx
        bericht.treffer.append(e)
        if laufendes_budget > token_budget:
            # The single, highest-ranked candidate alone exceeded the
            # budget -- it is still returned (an empty result would be
            # worse), but nothing further is added.
            LOGGER.warning(
                "docs_ingest: einzelner Treffer %s ueberschreitet bereits das Token-Budget (~%d)",
                e.pfad,
                token_budget,
            )
            for rest in kandidaten[kandidaten.index(e) + 1:]:
                grund = f"ueber dem Token-Budget (~{token_budget}) -- nach Relevanz verworfen"
                bericht.verworfen.append({"pfad": rest.pfad, "grund": grund})
                LOGGER.warning("docs_ingest: %s -- %s", rest.pfad, grund)
            break

    return bericht
