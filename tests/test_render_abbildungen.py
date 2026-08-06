"""Regression coverage for FINDINGS.md V-87 / BACKLOG E5-13.

`split_abbildungen()` had exactly one call site per renderer (inside the `kompetenzbezug`
verbatim path), so any other block carrying an `abbildungen` array was inert -- its
`(ABB:datei)` tokens printed as literal text and no image ever reached `word/media/` or an
HTML `<img>` tag. This file locks in the fix at three levels:

  1. the real E6-06 fixture's `beobachtungsbogen` `list` block, end to end through the CLI
     (docx zip + html string) -- the exact case FINDINGS.md V-87 measured;
  2. direct, block-type-generic checks (paragraph/callout/cards/table/fill_table/
     source_card/list) so the fix is not a `list`-only patch, per the task's instruction to
     route "any block carrying an abbildungen array" through the shared helper;
  3. the two safety invariants `split_abbildungen()`'s own docstring commits to: a block
     with no `abbildungen` renders markdown exactly as before, and a token with no matching
     metadata entry is kept as literal text, never silently dropped.
"""
from __future__ import annotations

import json
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

pytest.importorskip("docx", reason="python-docx is optional; DOCX abbildungen tests are skipped")
from docx import Document  # noqa: E402


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = REPO_ROOT / "plugin" / "skills" / "at-unterrichtsplanung" / "scripts"
RENDER_DOCUMENTS = SCRIPTS / "render_documents.py"
DOCX_CLI = SCRIPTS / "render_lesson_docx.py"
HTML_CLI = SCRIPTS / "render_lesson_html.py"
FIXTURE = (REPO_ROOT / "tests" / "fixtures" / "planning_flow" /
           "sek1_mathematik_k2_bruchzahlen.lesson.json")

# Real, shipped RIS formula images (V-87) -- reused rather than synthesized, so the test
# exercises actual add_picture()/base64 decode against real PNG bytes, not a placeholder.
IMG1_REL = "data/abbildungen/NOR40271471/hauptdokument.img8is.png"
IMG2_REL = "data/abbildungen/NOR40271471/hauptdokument.img9is.png"
IMG1_BYTES = (REPO_ROOT / "plugin" / IMG1_REL).read_bytes()
IMG2_BYTES = (REPO_ROOT / "plugin" / IMG2_REL).read_bytes()


def _run(cli: Path, src: Path, outfile: Path) -> None:
    result = subprocess.run(
        [sys.executable, str(cli), str(src), "-o", str(outfile)],
        cwd=REPO_ROOT, capture_output=True, text=True,
    )
    assert result.returncode == 0, result.stderr or result.stdout


# ---------------------------------------------------------------------------
# 1. The real fixture, end to end -- the exact case V-87 measured.
# ---------------------------------------------------------------------------

def test_e6_06_beobachtungsbogen_has_no_literal_tokens_and_embeds_both_images(tmp_path):
    result = subprocess.run(
        [sys.executable, str(RENDER_DOCUMENTS), str(FIXTURE), "--format", "both",
         "--outdir", str(tmp_path)],
        cwd=REPO_ROOT, capture_output=True, text=True,
    )
    assert result.returncode == 0, result.stderr or result.stdout

    fixture_data = json.loads(FIXTURE.read_text(encoding="utf-8"))
    beob = next(d for d in fixture_data["documents"] if d["id"] == "beobachtungsbogen")
    tokens = [
        meta["token"]
        for section in beob["sections"]
        for block in section["blocks"]
        for meta in (block.get("abbildungen") or [])
    ]
    assert len(tokens) == 2, "sanity: fixture must still carry exactly the two known tokens"

    with zipfile.ZipFile(tmp_path / "beobachtungsbogen.docx") as z:
        media_names = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media_names) == 2, f"expected 2 embedded images, found {media_names}"
        media_bytes = {z.read(n) for n in media_names}
        doc_xml = z.read("word/document.xml").decode("utf-8")
    assert media_bytes == {IMG1_BYTES, IMG2_BYTES}, (
        "embedded media must be the real shipped PNG bytes, not placeholders or corrupted"
    )
    for tok in tokens:
        assert tok not in doc_xml, f"literal token {tok!r} still present in word/document.xml"

    html = (tmp_path / "beobachtungsbogen.html").read_text(encoding="utf-8")
    for tok in tokens:
        assert tok not in html, f"literal token {tok!r} still present in beobachtungsbogen.html"
    assert html.count('<img class="abb"') == 2


# ---------------------------------------------------------------------------
# 2. Block-type-generic: not a `list`-only patch.
# ---------------------------------------------------------------------------

BLOCK_ABB = [
    {"token": "⟦ABB:one⟧", "datei": "img8is.png", "pfad": IMG1_REL},
    {"token": "⟦ABB:two⟧", "datei": "img9is.png", "pfad": IMG2_REL},
]

GENERIC_BLOCKS = {
    "paragraph": {"type": "paragraph", "text": "vor ⟦ABB:one⟧ nach ⟦ABB:two⟧ Ende",
                  "abbildungen": BLOCK_ABB},
    "list": {"type": "list", "label": "Punkte",
             "items": ["erster ⟦ABB:one⟧ Punkt", "zweiter ⟦ABB:two⟧ Punkt"],
             "abbildungen": BLOCK_ABB},
    "callout": {"type": "callout", "kind": "info", "label": "Hinweis",
                "text": "vor ⟦ABB:one⟧ nach ⟦ABB:two⟧ Ende", "abbildungen": BLOCK_ABB},
    "cards": {"type": "cards",
              "items": [{"title": "Karte", "text": "⟦ABB:one⟧ und ⟦ABB:two⟧"}],
              "abbildungen": BLOCK_ABB},
    "table": {"type": "table", "headers": ["A"],
              "rows": [["⟦ABB:one⟧"], ["⟦ABB:two⟧"]], "abbildungen": BLOCK_ABB},
    "fill_table": {"type": "fill_table", "headers": ["A"],
                   "rows": [["⟦ABB:one⟧"], ["⟦ABB:two⟧"]], "abbildungen": BLOCK_ABB},
    "source_card": {"type": "source_card", "title": "Quelle",
                    "excerpt": "⟦ABB:one⟧ und ⟦ABB:two⟧", "abbildungen": BLOCK_ABB},
}


def _doc_for(block: dict) -> dict:
    return {"title": "T", "sections": [{"heading": "S", "blocks": [block]}]}


@pytest.mark.parametrize("name", sorted(GENERIC_BLOCKS))
def test_block_type_with_abbildungen_embeds_images_not_literal_tokens(name, tmp_path):
    block = GENERIC_BLOCKS[name]
    src = tmp_path / f"{name}.lesson.json"
    src.write_text(json.dumps(_doc_for(block)), encoding="utf-8")

    docx_out = tmp_path / f"{name}.docx"
    _run(DOCX_CLI, src, docx_out)
    with zipfile.ZipFile(docx_out) as z:
        media_names = [n for n in z.namelist() if n.startswith("word/media/")]
        media_bytes = {z.read(n) for n in media_names}
        doc_xml = z.read("word/document.xml").decode("utf-8")
    assert len(media_names) == 2, f"{name}: expected 2 embedded images, found {media_names}"
    assert media_bytes == {IMG1_BYTES, IMG2_BYTES}, f"{name}: embedded bytes do not match source"
    assert "⟦ABB:one⟧" not in doc_xml and "⟦ABB:two⟧" not in doc_xml, (
        f"{name}: a literal token survived into word/document.xml"
    )

    html_out = tmp_path / f"{name}.html"
    _run(HTML_CLI, src, html_out)
    html = html_out.read_text(encoding="utf-8")
    assert html.count('<img class="abb"') == 2, f"{name}: expected 2 <img class=\"abb\"> tags"
    assert "⟦ABB:one⟧" not in html and "⟦ABB:two⟧" not in html, (
        f"{name}: a literal token survived into the html"
    )


# ---------------------------------------------------------------------------
# 3a. Safety invariant: an unmatched token is kept as literal text, never dropped.
# ---------------------------------------------------------------------------

def test_unmatched_token_is_kept_literal_not_silently_dropped(tmp_path):
    block = {
        "type": "paragraph",
        "text": "vor ⟦ABB:missing⟧ nach",
        # abbildungen is present (non-empty) but has no entry for the token actually in
        # `text` -- split_abbildungen()'s documented safe-failure mode.
        "abbildungen": [{"token": "⟦ABB:other⟧", "datei": "x.png", "pfad": IMG1_REL}],
    }
    src = tmp_path / "unmatched.lesson.json"
    src.write_text(json.dumps(_doc_for(block)), encoding="utf-8")

    docx_out = tmp_path / "unmatched.docx"
    _run(DOCX_CLI, src, docx_out)
    with zipfile.ZipFile(docx_out) as z:
        media_names = [n for n in z.namelist() if n.startswith("word/media/")]
        doc_xml = z.read("word/document.xml").decode("utf-8")
    assert media_names == [], "no image should be embedded for an unmatched token"
    assert "⟦ABB:missing⟧" in doc_xml, (
        "an unmatched token must survive as visible literal text, not vanish"
    )

    html_out = tmp_path / "unmatched.html"
    _run(HTML_CLI, src, html_out)
    html = html_out.read_text(encoding="utf-8")
    assert '<img class="abb"' not in html
    assert "⟦ABB:missing⟧" in html, (
        "an unmatched token must survive as visible literal text, not vanish"
    )


# ---------------------------------------------------------------------------
# 3b. Safety invariant: a block with no `abbildungen` still renders markdown normally
#     (add_md_abb()/md_abb() must fall through to plain add_md()/md() unchanged).
# ---------------------------------------------------------------------------

def test_block_without_abbildungen_still_renders_markdown_normally(tmp_path):
    block = {"type": "paragraph", "text": "**fett** und *kursiv*"}  # no "abbildungen" key
    src = tmp_path / "plain.lesson.json"
    src.write_text(json.dumps(_doc_for(block)), encoding="utf-8")

    docx_out = tmp_path / "plain.docx"
    _run(DOCX_CLI, src, docx_out)
    doc = Document(str(docx_out))
    runs = [r for p in doc.paragraphs for r in p.runs if r.text]
    bold_texts = [r.text for r in runs if r.bold]
    italic_texts = [r.text for r in runs if r.italic]
    assert "fett" in bold_texts, "markdown bold must still parse with no abbildungen present"
    assert "kursiv" in italic_texts, "markdown italic must still parse with no abbildungen present"

    html_out = tmp_path / "plain.html"
    _run(HTML_CLI, src, html_out)
    html = html_out.read_text(encoding="utf-8")
    assert "<b>fett</b>" in html
    assert "<i>kursiv</i>" in html
